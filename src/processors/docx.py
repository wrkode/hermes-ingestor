"""
Microsoft Word (DOCX) document processor.
"""

import os
from typing import Dict, Any
import docx

from .base import BaseProcessor


class DocxProcessor(BaseProcessor):
    """Processor for Microsoft Word (DOCX) documents."""

    def extract_text(self) -> str:
        """
        Extract text content from a DOCX file.

        Returns:
            str: The extracted text
        """
        doc = docx.Document(self.file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)
        
        return "\n".join(full_text)

    def extract_metadata(self) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX file.

        Returns:
            dict: DOCX metadata including title, author, etc.
        """
        metadata = {
            "source": self.file_path,
            "filename": os.path.basename(self.file_path),
            "file_type": "docx",
        }

        doc = docx.Document(self.file_path)
        
        # Extract core properties if available
        if hasattr(doc, "core_properties"):
            props = doc.core_properties
            
            if hasattr(props, "title") and props.title:
                metadata["title"] = props.title
                
            if hasattr(props, "author") and props.author:
                metadata["author"] = props.author
                
            if hasattr(props, "created") and props.created:
                metadata["created"] = props.created
                
            if hasattr(props, "modified") and props.modified:
                metadata["modified"] = props.modified
                
            if hasattr(props, "last_modified_by") and props.last_modified_by:
                metadata["last_modified_by"] = props.last_modified_by
                
            if hasattr(props, "category") and props.category:
                metadata["category"] = props.category
                
            if hasattr(props, "comments") and props.comments:
                metadata["comments"] = props.comments
                
            if hasattr(props, "keywords") and props.keywords:
                metadata["keywords"] = props.keywords
        
        # Count paragraphs, tables, and approximate word count
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        
        # Get approximate word count
        text = self.extract_text()
        metadata["word_count"] = len(text.split())

        return metadata 